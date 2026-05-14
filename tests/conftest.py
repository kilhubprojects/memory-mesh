"""Shared pytest fixtures and configuration for MemoryMesh tests."""

from __future__ import annotations

from pathlib import Path

import pytest


@pytest.fixture()
def tmp_db_path(tmp_path: Path) -> Path:
    """Return a fresh temporary directory for database files."""
    db_dir = tmp_path / "db"
    db_dir.mkdir()
    return db_dir


@pytest.fixture()
def sample_text_file(tmp_path: Path) -> Path:
    """A plain-text file with two paragraphs."""
    p = tmp_path / "sample.txt"
    p.write_text(
        "Hello world.\n\nThis is the second paragraph with more content to read.",
        encoding="utf-8",
    )
    return p


@pytest.fixture()
def sample_py_file(tmp_path: Path) -> Path:
    """A minimal Python source file with one function and one class."""
    p = tmp_path / "sample.py"
    py_src = (
        'def greet(name: str) -> str:\n    """Return a greeting."""\n'
        '    return f"Hello, {name}!"\n\n\n'
        'class Greeter:\n    def hello(self) -> str:\n        return "hi"\n'
    )
    p.write_text(py_src, encoding="utf-8")
    return p


@pytest.fixture()
def sample_markdown_file(tmp_path: Path) -> Path:
    """A Markdown file with two headings."""
    p = tmp_path / "sample.md"
    md_src = (
        "# Introduction\n\nThis is the intro section.\n\n"
        "## Details\n\nMore detailed content here.\n"
    )
    p.write_text(md_src, encoding="utf-8")
    return p


@pytest.fixture()
def sample_pdf_path(tmp_path: Path) -> Path:
    """A minimal single-page PDF generated with fpdf2."""
    try:
        from fpdf import FPDF  # type: ignore[import-untyped]
    except ImportError:
        pytest.skip("fpdf2 not installed — skipping PDF fixture")

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.cell(text="MemoryMesh PDF test fixture")
    p = tmp_path / "sample.pdf"
    pdf.output(str(p))
    return p


@pytest.fixture()
def sample_docx_path(tmp_path: Path) -> Path:
    """A minimal .docx file generated with python-docx."""
    try:
        import docx  # type: ignore[import-untyped]
    except ImportError:
        pytest.skip("python-docx not installed — skipping DOCX fixture")

    document = docx.Document()
    document.add_paragraph("MemoryMesh DOCX test fixture")
    document.add_paragraph("Second paragraph for testing.")
    p = tmp_path / "sample.docx"
    document.save(str(p))
    return p

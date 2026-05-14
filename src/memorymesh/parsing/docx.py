"""DOCX parser backed by python-docx.

Extracts paragraph text in document order.  Tables and headers/footers are
currently skipped (MVP scope).  The daemon never crashes on a malformed DOCX.
"""

from __future__ import annotations

from pathlib import Path

from loguru import logger

from memorymesh.core.models import ParsedDocument
from memorymesh.parsing.base import Parser


class DocxParser(Parser):
    """Parser for Microsoft Word .docx files."""

    @property
    def supported_extensions(self) -> frozenset[str]:
        """File extensions handled by this parser: Microsoft Word documents."""
        return frozenset({".docx"})

    def parse(self, path: Path) -> ParsedDocument:
        """Extract paragraph text from a .docx file.

        Args:
            path: Absolute path to the .docx file.

        Returns:
            :class:`~memorymesh.core.models.ParsedDocument` with paragraph text joined by newlines.
        """
        try:
            import docx  # python-docx — deferred import
        except ImportError:
            msg = "python-docx is not installed; run 'uv add python-docx'"
            logger.error(msg)
            return ParsedDocument(path=path, text="", file_type=".docx", metadata={"error": msg})

        meta: dict[str, object] = {}
        try:
            doc = docx.Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n".join(paragraphs)
            meta["paragraph_count"] = len(paragraphs)
            logger.debug(f"DocxParser: {path.name!r} — {len(paragraphs)} paragraphs")
        except Exception as exc:
            msg = f"python-docx failed to read {path}: {exc}"
            logger.warning(msg)
            meta["error"] = msg
            return ParsedDocument(path=path, text="", file_type=".docx", metadata=meta)

        return ParsedDocument(path=path, text=text, file_type=".docx", metadata=meta)
